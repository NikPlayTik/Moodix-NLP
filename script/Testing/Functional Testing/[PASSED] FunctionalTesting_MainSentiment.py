# -*- coding: utf-8 -*-
"""
Auto-reporting pytest hook:
    * Collects duration/outcome per test.
    * At session end, builds Word report with summary table and embedded bar chart.
    * Saves only the Word report to the specified REPORT_DIR.
"""
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
import datetime, os

# Directory where the final Word report will be saved
REPORT_DIR = r"D:/Laboratary and practical/!Диплом/Model/Отчеты о тестировании"

def pytest_configure(config):
    config._test_results = []

def pytest_runtest_logreport(report):
    if report.when == 'call':
        config = report.config
        config._test_results.append({
            'id': report.nodeid,
            'outcome': report.outcome,
            'duration': report.duration,
        })

def pytest_sessionfinish(session, exitstatus):
    results = session.config._test_results
    if not results:
        return

    # Ensure the report directory exists
    os.makedirs(REPORT_DIR, exist_ok=True)

    # Sort and prepare chart data
    results_sorted = sorted(results, key=lambda x: x['duration'], reverse=True)
    labels = [r['id'].split("::")[-1] for r in results_sorted]
    durations = [r['duration'] for r in results_sorted]

    # Plot horizontal bar chart
    plt.figure(figsize=(10, max(4, len(labels)*0.3)))
    plt.barh(labels, durations)
    plt.xlabel('Duration (s)')
    plt.title('Test Durations')
    plt.tight_layout()
    chart_path = os.path.join(REPORT_DIR, 'test_durations.png')
    plt.savefig(chart_path, dpi=150)
    plt.close()

    # Build Word document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0,0,0)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.add_paragraph(f'Отчёт автотестов — {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}')

    # Summary table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Тест'
    hdr[1].text = 'Статус'
    hdr[2].text = 'Время, c'

    for r in results:
        row = table.add_row().cells
        row[0].text = r['id']
        row[1].text = 'PASSED' if r['outcome']=='passed' else r['outcome'].upper()
        row[2].text = f"{r['duration']:.3f}"

    doc.add_paragraph()
    doc.add_picture(chart_path, width=Inches(6))

    report_name = f"Test_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    report_path = os.path.join(REPORT_DIR, report_name)
    doc.save(report_path)

    # Cleanup: remove temporary chart file
    try:
        os.remove(chart_path)
    except OSError:
        pass

    print(f"\nWord‑отчёт сохранён: {report_path}\n")

"""
2_FunctionalTesting: Functional / integration tests for SentimentAnalyzer with auto Word reporting.

Запуск:
    pytest 2_FunctionalTesting_MainSentiment.py

Перед запуском:
    set MOODIX_MODEL=D:\Laboratary and practical\!Диплом\Model\14 Moodix v0.64\model.keras
    set MOODIX_TOKENIZER=D:\Laboratary and practical\!Диплом\Model\14 Moodix v0.64\tokenizer.pickle
"""
import os, time, pytest
from Start_AVG import SentimentAnalyzer

MODEL_PATH = os.getenv("MOODIX_MODEL", "model.keras")
TOKENIZER_PATH = os.getenv("MOODIX_TOKENIZER", "tokenizer.pickle")

@pytest.fixture(scope="session")
def analyzer():
    sa = SentimentAnalyzer(MODEL_PATH, TOKENIZER_PATH)
    sa.load_model()
    return sa

# ---------------- Типовые позитив/негатив/нейтрал ---------------
cases = [
    ("Смартфон Samsung Galaxy A55 5G: переоцененный представитель А-серии", "neutrally"),
    ("К огромному сожалению, для моих глаз обзор новых смартфонов на строке «сенсорный дисплей AMOLED» заканчивается. :(", "neutrally"),
    ("Что за левый браузер? У меня в Firefox JetStream2 выдаёт 66,7. Octane 2.0 выдаёт 19020. И это базовая Pura 70 с Kirin 9000S1, да без режима максимальной производительности.", "neutrally"),
    ("В Беларуси достроили самую высокую панельку. Нет, не в Минске", "neutrally"),
    ("Слабоватый аккумулятор, раз в сутки приходится заряжать в любом случае, даже без использования нагруженных приложений. Еще небольшой минус в расположении дополнительных кнопок на корпусе, польза от них минимальная, можно было обыграть получше их возможности", "neutrally"),
    ("Очень удобный в пользовании. Камера и экран супер. Батарея по сравнению с СЕ 2020 кажется не садится.", "positive"),
    ("Приобрела айфон 16 буквально пару дней назад. Ранее пользовалась 13 мини. Что хочу сказать? Этот телефон шикарен. Да, это не про версия, тут нет третьей камеры и 120 Гц. Но если смотреть правде в глаза, это не всем и нужно. Мне все кричали в три голоса взять про-шку, но меня зацепила 16-тка бирюзовая. Телефон мощный, камеры усовершенствовали, появились новые фишки с новыми кнопками, динамический остров добавляет комфорта в использовании. Мне все нравится, очень удобный и классный телефон. Думала перейти с айфона на андроид, но пришла к выводу, что даже суперфлагманы работают хуже базовых новых айфонов, а стоят намного дороже.", "positive"),
    ("Ужасный телефон не стоит своих денег. Достоинства: их в принципе нет. Недостатки: очень быстро нагревается процессор при обычном использовании", "negative"),
    ("Сильно разочарован этим устройством на фоне даже того же SGS23U. Радости не было предела, когда наконец довели до ума частоту ШИМ, вместе с этим были надежды, что исправили проблему отображения белого при наклонах экрана, то есть углы обзора. Начитавшись хвалебных обзоров наконец ломанулся в магазин что бы купить это великолепное устройство по всем обзорам и по многим отзывам, но что в итоге? Спасибо за частоту ШИМ, спасибо за поляризационную пленку, но углы обзора стали ещё хуже, чем были во всех флагманах этой линейки, теперь белый уходит в синий при самом малом отклонении устройства в сторону, при большем уходит в радугу, кроме этого очень неравномерная заливка белого на экране, даже если смотреть на него прямо в упор. Казалось бы, Самсунг впихнул в свое детище передовой экран новейшего поколения, продает свои экраны другим производителям, в которых всех этих проблем нет, что же с самсунгом не так? Конуренты с экранами ВОЕ предлагают сегодня практически идеальный баланс белого не в зависимости от угла обзора и при этом гибкую настройку цветопередачи на вкус каждого, в т.ч. и бб. Ярким экран тоже не назвать, на 'экстремальная яркость' в прямом сравнении не ярче моего старенького 1+9R, который вышел 3.5 года назад. Без этого костыля и вовсе тусклее. У друга флагманский смартфон из концерта bbk, от экрана глаз не оторвать( А в помещении экран светит, как прожектор. Стилус порезали. Дизайн удешевили, в руках флагман теперь как китайская побрекушка. Очень обидно наблюдать за этой линейкой, которая теряет свой фарм и изюминку, компромиссы, которых с каждым годом все больше и больше", "negative"),
    ("Ужасно! Только зарегистрировался и телефон получил удалённую блокировку UK Samsung Electronics. Стандартно все настраивал, все галочки снимал.", "negative"),
    ("Отличный аппарат. Очень компактный, удобный. Оболочка лучшая. Камера, звук супер.", "positive"),
    ("Перешла с Samsung S10plus. Разницы, кроме того, что батарея лучше, чем у 6-летнего самсунга нет вообще, к тому же на iphone15 не работает дублирование экрана... Не советую, потратьте деньги на чо-нить другое", "negative")
]

@pytest.mark.parametrize("text,expected", cases)
def test_main_sentiment(analyzer, text, expected):
    res = analyzer.analyze_sentiment(text)
    assert res["main"]["label"] == expected
    assert res["main"]["confidence"] > 0.50

# ---------------- SLA (<1 s) ------------------------------------
def test_single_inference_time(analyzer):
    start = time.time()
    analyzer.analyze_sentiment("Быстрая проверка производительности.")
    assert time.time() - start < 1.0
